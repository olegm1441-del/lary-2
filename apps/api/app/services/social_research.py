from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable, Literal
from urllib.parse import urlparse

import httpx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pydantic import BaseModel, Field, field_validator

from app.services import ai_router


USER_FRIENDLY_SOURCE_ERROR = "Не удалось получить проверяемые источники. Черновик сохранён. Попробуйте повторить позже."
USER_FRIENDLY_AI_ERROR = "Не удалось подготовить доказательный материал. Черновик сохранён. Попробуйте повторить позже."
SOURCE_REQUEST_HEADERS = {"User-Agent": "LARI/0.1 (legacyinfo@yandex.ru)"}
TRUSTED_SOURCE_HOSTS_WITH_BROKEN_CERTIFICATE_CHAIN = {"66.rosstat.gov.ru"}


class SocialResearchGenerationError(ValueError):
    pass


class VerifiedSource(BaseModel):
    source_id: str = Field(..., min_length=3, max_length=80)
    title: str = Field(..., min_length=5, max_length=240)
    publisher: str = Field(..., min_length=2, max_length=180)
    publication_date: str
    url: str
    source_type: Literal["official", "research", "regional", "sector"]
    territory: str = Field(..., min_length=2, max_length=180)
    claim: str = Field(..., min_length=10, max_length=900)
    evidence: str = Field(..., min_length=10, max_length=1200)
    verified_at: str

    @field_validator("publication_date")
    @classmethod
    def publication_is_current(cls, value: str) -> str:
        match = re.fullmatch(r"(\d{4})(?:-\d{2}-\d{2})?", value)
        if not match or int(match.group(1)) < 2024:
            raise ValueError("Источник должен быть опубликован не ранее 2024 года.")
        return value

    @field_validator("url")
    @classmethod
    def url_is_public_http(cls, value: str) -> str:
        parsed = urlparse(value)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            raise ValueError("Источник должен иметь публичный HTTP(S) URL.")
        return value


class EvidenceItem(BaseModel):
    claim: str = Field(..., min_length=10, max_length=700)
    source_id: str = Field(..., min_length=3, max_length=80)
    application: str = Field(..., min_length=10, max_length=500)


class SurveyBlock(BaseModel):
    hypothesis: str = Field(..., min_length=10, max_length=500)
    question: str = Field(..., min_length=10, max_length=400)
    answers: list[str] = Field(..., min_length=4, max_length=6)
    interpretation: str = Field(..., min_length=10, max_length=500)

    @field_validator("answers")
    @classmethod
    def answers_are_distinct(cls, answers: list[str]) -> list[str]:
        normalized = [re.sub(r"\s+", " ", answer).strip() for answer in answers if answer.strip()]
        if len(normalized) < 4 or len(normalized) > 6 or len({item.casefold() for item in normalized}) != len(normalized):
            raise ValueError("Нужно от 4 до 6 разных вариантов ответа.")
        return normalized


class SocialResearchOutput(BaseModel):
    document_title: str = Field(..., min_length=10, max_length=180)
    initial_situation: str = Field(..., min_length=20, max_length=1600)
    target_group: str = Field(..., min_length=5, max_length=700)
    social_problem: str = Field(..., min_length=20, max_length=1400)
    significance: str = Field(..., min_length=20, max_length=1600)
    project_response: str = Field(..., min_length=20, max_length=1400)
    evidence_items: list[EvidenceItem] = Field(..., min_length=2, max_length=8)
    survey: SurveyBlock
    useful_material_ids: list[str] = Field(default_factory=list, max_length=8)
    manual_checks: list[str] = Field(default_factory=list, max_length=6)


@dataclass(frozen=True)
class SocialResearchDocument:
    output_path: Path
    output: SocialResearchOutput
    sources: list[VerifiedSource]
    sections: list[dict[str, str]]


SOURCE_CATALOG = [
    {
        "source_id": "sverdlovskstat-culture-2024",
        "title": "Свердловская область. Культура, туризм и отдых: театры",
        "publisher": "Свердловскстат",
        "publication_date": "2026",
        "url": "https://66.rosstat.gov.ru/storage/mediabank/04001-2024-%D1%81%D0%B0%D0%B9%D1%82.pdf",
        "source_type": "regional",
        "territory": "Свердловская область",
        "claim": "В 2024 году в Свердловской области работали 25 профессиональных театров, которые провели 8 072 мероприятия; число зрителей составило 1 430,4 тыс. человек.",
        "evidence": "Официальный статистический сборник содержит региональные показатели театральной инфраструктуры и посещаемости за 2024 год.",
        "tags": ["театр", "театраль", "культур", "искусств"],
    },
    {
        "source_id": "wciom-theatre-2024",
        "title": "Весь мир — театр!",
        "publisher": "ВЦИОМ",
        "publication_date": "2024-02-26",
        "url": "https://wciom.ru/analytical-reviews/analiticheskii-obzor/ves-mir-teatr",
        "source_type": "research",
        "territory": "Российская Федерация",
        "claim": "По опросу ВЦИОМ 2024 года, 25% россиян посещают театр несколько раз в год и чаще; среди молодёжи 18–24 лет этот показатель составляет 39%.",
        "evidence": "Материал содержит результаты всероссийского телефонного опроса 1 600 россиян, включая распределения по возрасту и типу населённого пункта.",
        "tags": ["театр", "театраль", "культур", "искусств", "молод"],
    },
]


def validate_verified_sources(sources: list[VerifiedSource]) -> list[VerifiedSource]:
    seen_ids: set[str] = set()
    seen_urls: set[str] = set()
    for source in sources:
        normalized_url = source.url.rstrip("/").casefold()
        if source.source_id in seen_ids or normalized_url in seen_urls:
            raise ValueError("Один источник не должен дублироваться.")
        seen_ids.add(source.source_id)
        seen_urls.add(normalized_url)
    return sources


def retrieve_verified_sources(inputs: dict) -> list[VerifiedSource]:
    query = " ".join(
        str(inputs.get(key) or "") for key in ("direction", "problem", "target_group", "project_response", "region")
    ).casefold()
    region = str(inputs.get("region") or "").casefold().replace("ё", "е")
    candidates = [
        item
        for item in SOURCE_CATALOG
        if any(tag.casefold() in query for tag in item["tags"])
        and (
            item["territory"] == "Российская Федерация"
            or item["territory"].casefold().replace("ё", "е") in region
        )
    ]
    verified: list[VerifiedSource] = []
    now = datetime.now(timezone.utc).isoformat()
    with httpx.Client(follow_redirects=True, timeout=10.0, headers=SOURCE_REQUEST_HEADERS) as client:
        for item in candidates:
            try:
                response = _fetch_public_source(client, item["url"])
            except httpx.HTTPError:
                continue
            if response.status_code != 200:
                continue
            final_url = str(response.url)
            if "search" in urlparse(final_url).path.casefold():
                continue
            verified.append(
                VerifiedSource.model_validate(
                    {
                        key: value
                        for key, value in item.items()
                        if key != "tags"
                    }
                    | {"url": final_url, "verified_at": now}
                )
            )
        required_types_present = (
            any(source.source_type in {"official", "regional"} for source in verified)
            and any(source.source_type == "research" for source in verified)
        )
        if not required_types_present:
            for candidate in _discover_source_candidates(inputs):
                verified_source = _verify_discovered_source(client, candidate, now)
                if verified_source and all(
                    item.source_id != verified_source.source_id
                    and item.url.rstrip("/").casefold() != verified_source.url.rstrip("/").casefold()
                    for item in verified
                ):
                    verified.append(verified_source)
    return validate_verified_sources(verified)


def _fetch_public_source(client: httpx.Client, url: str) -> httpx.Response:
    """Fetch a source normally, with a narrow fallback for a known Rosstat TLS-chain defect."""
    try:
        return client.get(url)
    except httpx.ConnectError as exc:
        host = (urlparse(url).hostname or "").casefold()
        certificate_failure = "certificate_verify_failed" in str(exc).casefold()
        if host not in TRUSTED_SOURCE_HOSTS_WITH_BROKEN_CERTIFICATE_CHAIN or not certificate_failure:
            raise
        with httpx.Client(
            follow_redirects=True,
            timeout=10.0,
            headers=SOURCE_REQUEST_HEADERS,
            verify=False,
        ) as fallback_client:
            return fallback_client.get(url)


def _discover_source_candidates(inputs: dict) -> list[dict]:
    prompt = (
        "Найди до 6 прямых публичных источников для доказательной базы грантового проекта. "
        "Нужны публикации не ранее 2024 года: минимум один официальный федеральный или региональный источник "
        "и минимум одно исследование ВЦИОМ, ФОМ, НИУ ВШЭ или другой исследовательской организации. "
        "Не давай страницы поиска, главные страницы и вымышленные ссылки. Не пересказывай источник шире опубликованных данных. "
        "Верни только JSON без markdown: "
        '{"sources":[{"source_id":"short-id","title":"...","publisher":"...",'
        '"publication_date":"YYYY-MM-DD","url":"https://...","source_type":"official|regional|research|sector",'
        '"territory":"...","claim":"проверяемое утверждение","evidence":"что именно опубликовано"}]}. '
        f"Данные проекта: {json.dumps(inputs, ensure_ascii=False, sort_keys=True)}"
    )
    try:
        raw = ai_router.generate_with_gigachat(prompt)
        parsed = json.loads(_extract_json_object_text(raw))
    except Exception:
        return []
    sources = parsed.get("sources") if isinstance(parsed, dict) else None
    return [item for item in sources[:6] if isinstance(item, dict)] if isinstance(sources, list) else []


def _verify_discovered_source(
    client: httpx.Client,
    candidate: dict,
    verified_at: str,
) -> VerifiedSource | None:
    try:
        provisional = VerifiedSource.model_validate(candidate | {"verified_at": verified_at})
        response = _fetch_public_source(client, provisional.url)
        if response.status_code != 200:
            return None
        final_url = str(response.url)
        parsed = urlparse(final_url)
        if "search" in parsed.path.casefold() or len(response.content) < 500:
            return None
        page_text = re.sub(r"\s+", " ", response.text).casefold()
        claim_numbers = _numbers(provisional.claim)
        normalized_page = re.sub(r"\s+", "", page_text)
        if claim_numbers and any(number.replace(",", ".") not in normalized_page.replace(",", ".") for number in claim_numbers):
            return None
        evidence_tokens = {
            token
            for token in re.findall(r"[а-яёa-z]{5,}", f"{provisional.claim} {provisional.evidence}".casefold())
            if token not in {"который", "данные", "публикация", "материал", "источник"}
        }
        page_tokens = set(re.findall(r"[а-яёa-z]{5,}", page_text))
        if len(evidence_tokens.intersection(page_tokens)) < min(3, len(evidence_tokens)):
            return None
        return provisional.model_copy(update={"url": final_url})
    except (ValueError, httpx.HTTPError):
        return None


def validate_social_research_output(
    output: SocialResearchOutput,
    sources: list[VerifiedSource],
    inputs: dict,
) -> SocialResearchOutput:
    source_by_id = {source.source_id: source for source in validate_verified_sources(sources)}
    referenced = [item.source_id for item in output.evidence_items]
    unknown = sorted(set(referenced + output.useful_material_ids) - set(source_by_id))
    if unknown:
        raise ValueError(f"AI использовал неизвестный source_id: {', '.join(unknown)}")

    output_text = " ".join(
        [
            output.document_title,
            output.initial_situation,
            output.target_group,
            output.social_problem,
            output.significance,
            output.project_response,
            *(item.claim for item in output.evidence_items),
            *(item.application for item in output.evidence_items),
            output.survey.hypothesis,
            output.survey.question,
            *output.survey.answers,
            output.survey.interpretation,
            *output.manual_checks,
        ]
    )
    allowed_text = " ".join(
        [
            *(str(value) for value in inputs.values()),
            *(source.claim for source in sources),
            *(source.evidence for source in sources),
            *(source.publication_date for source in sources),
        ]
    )
    unsupported_numbers = _numbers(output_text) - _numbers(allowed_text)
    if unsupported_numbers:
        raise ValueError(f"Найдены неподтверждённые числа: {', '.join(sorted(unsupported_numbers))}")

    allowed_locations = _locations(allowed_text)
    invented_locations = _locations(output_text) - allowed_locations
    if invented_locations:
        raise ValueError(f"Найдена неподтверждённая география: {', '.join(sorted(invented_locations))}")
    input_text = " ".join(str(value) for value in inputs.values())
    for item in output.evidence_items:
        source = source_by_id[item.source_id]
        item_allowed = f"{input_text} {source.claim} {source.evidence} {source.publication_date}"
        unsupported_item_numbers = _numbers(item.claim) - _numbers(item_allowed)
        if unsupported_item_numbers:
            raise ValueError(
                f"Факт содержит число не из указанного источника: {', '.join(sorted(unsupported_item_numbers))}"
            )
    return output


def build_social_research_document(
    inputs: dict,
    *,
    output_path: Path | None = None,
    source_fetcher: Callable[[dict], list[VerifiedSource]] | None = None,
    ai_generator: Callable[[str], str] | None = None,
) -> SocialResearchDocument:
    if os.getenv("APP_ENV") == "test" and source_fetcher is None and ai_generator is None:
        source_fetcher = _test_verified_sources
        ai_generator = lambda _: json.dumps(_test_output(inputs), ensure_ascii=False)
    source_fetcher = source_fetcher or retrieve_verified_sources
    ai_generator = ai_generator or ai_router.generate_with_gigachat
    sources = validate_verified_sources(source_fetcher(inputs))
    if not sources or not any(source.source_type in {"official", "regional"} for source in sources):
        raise SocialResearchGenerationError(USER_FRIENDLY_SOURCE_ERROR)
    if not any(source.source_type == "research" for source in sources):
        raise SocialResearchGenerationError(USER_FRIENDLY_SOURCE_ERROR)

    prompt = _build_prompt(inputs, sources)
    error: Exception | None = None
    for attempt in range(2):
        try:
            raw = ai_generator(prompt if attempt == 0 else _repair_prompt(prompt, error))
            output = SocialResearchOutput.model_validate(json.loads(_extract_json_object_text(raw)))
            validate_social_research_output(output, sources, inputs)
            break
        except Exception as exc:  # provider/parser details stay internal
            error = exc
    else:
        raise SocialResearchGenerationError(USER_FRIENDLY_AI_ERROR) from error

    resolved_path = output_path or Path("social-research.docx")
    resolved_path.parent.mkdir(parents=True, exist_ok=True)
    _render_docx(resolved_path, output, sources, inputs)
    return SocialResearchDocument(
        output_path=resolved_path,
        output=output,
        sources=sources,
        sections=_result_sections(output, sources),
    )


def _test_verified_sources(_: dict) -> list[VerifiedSource]:
    now = datetime.now(timezone.utc).isoformat()
    return [
        VerifiedSource(
            source_id="test-official-2024",
            title="Тестовый официальный источник",
            publisher="Официальная организация",
            publication_date="2024-01-01",
            url="https://example.org/official-2024",
            source_type="official",
            territory="Российская Федерация",
            claim="Официальный материал содержит сведения по теме проекта.",
            evidence="Источник используется только в изолированных автоматических тестах.",
            verified_at=now,
        ),
        VerifiedSource(
            source_id="test-research-2024",
            title="Тестовый исследовательский источник",
            publisher="Исследовательская организация",
            publication_date="2024-02-01",
            url="https://example.org/research-2024",
            source_type="research",
            territory="Российская Федерация",
            claim="Исследовательский материал описывает особенности целевой группы.",
            evidence="Источник используется только в изолированных автоматических тестах.",
            verified_at=now,
        ),
    ]


def _test_output(inputs: dict) -> dict:
    direction = str(inputs.get("direction") or "проекта")
    region = str(inputs.get("region") or "указанная территория")
    target_group = str(inputs.get("target_group") or "целевая группа проекта")
    problem = str(inputs.get("problem") or "требуется уточнить исходную проблему проекта")
    response = str(
        inputs.get("project_response")
        or "Проект предлагает содержательную работу с обозначенной проблемой и выбранной целевой группой."
    )
    return {
        "document_title": f"Актуальность проекта по направлению «{direction}»",
        "initial_situation": f"Исходная ситуация описана пользователем для направления «{direction}» на территории «{region}» и требует доказательной проверки.",
        "target_group": target_group,
        "social_problem": f"Пользователь описывает проблему проекта следующим образом: {problem}.",
        "significance": "Обозначенная проблема влияет на доступ целевой группы к соответствующим возможностям участия.",
        "project_response": response,
        "evidence_items": [
            {
                "claim": "Официальный материал содержит сведения по теме проекта.",
                "source_id": "test-official-2024",
                "application": "Использовать для проверки исходной ситуации и территории проекта.",
            },
            {
                "claim": "Исследовательский материал описывает особенности целевой группы.",
                "source_id": "test-research-2024",
                "application": "Использовать для уточнения характеристик целевой группы.",
            },
        ],
        "survey": {
            "hypothesis": "Целевая группа сталкивается с обозначенной пользователем проблемой.",
            "question": "Насколько актуальна для вас описанная проблема?",
            "answers": ["Очень актуальна", "Скорее актуальна", "Скорее не актуальна", "Совсем не актуальна"],
            "interpretation": "Опрос ещё не проведён; ответы помогут проверить пользовательскую гипотезу.",
        },
        "useful_material_ids": ["test-official-2024", "test-research-2024"],
        "manual_checks": ["Заменить тестовые материалы реальными источниками при production-проверке."],
    }


def _build_prompt(inputs: dict, sources: list[VerifiedSource]) -> str:
    source_payload = [source.model_dump() for source in sources]
    return (
        "Ты профессиональный грантрайтер ПФКИ. Подготовь доказательный материал только на основе входных данных "
        "и immutable verified sources ниже. Не придумывай URL, числа, города, учреждения или проведённые опросы. "
        "Не меняй source_id. Пользовательская проблема — гипотеза проекта, а не уже доказанный факт. "
        "Верни строго JSON без markdown по схеме: "
        '{"document_title":"...","initial_situation":"...","target_group":"...",'
        '"social_problem":"...","significance":"...","project_response":"...",'
        '"evidence_items":[{"claim":"...","source_id":"source-1","application":"..."}],'
        '"survey":{"hypothesis":"...","question":"...","answers":["..."],"interpretation":"..."},'
        '"useful_material_ids":["source-1"],"manual_checks":["..."]}. '
        "В survey дай 4–6 взаимоисключающих ответов и ясно укажи, что опрос ещё не проведён. "
        f"Входные данные: {json.dumps(inputs, ensure_ascii=False, sort_keys=True)}. "
        f"Проверенные источники: {json.dumps(source_payload, ensure_ascii=False, sort_keys=True)}."
    )


def _repair_prompt(prompt: str, error: Exception | None) -> str:
    return (
        f"{prompt}\nПредыдущий ответ не прошёл проверку ({type(error).__name__ if error else 'validation'}). "
        "Исправь JSON, используй только известные source_id, числа и географию из входных данных или источников. "
        "Верни только JSON."
    )


def _extract_json_object_text(raw: str) -> str:
    text = str(raw or "").strip()
    fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, flags=re.IGNORECASE | re.DOTALL)
    if fenced:
        return fenced.group(1)
    start = text.find("{")
    end = text.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("Ответ не содержит JSON.")
    return text[start : end + 1]


def _result_sections(output: SocialResearchOutput, sources: list[VerifiedSource]) -> list[dict[str, str]]:
    source_by_id = {source.source_id: source for source in sources}
    evidence_lines = []
    for item in output.evidence_items:
        source = source_by_id[item.source_id]
        evidence_lines.append(
            f"{item.claim}\nИсточник: {source.publisher}, {source.publication_date}. {source.url}\nКак использовать: {item.application}"
        )
    return [
        {"title": "Исходная ситуация", "body": output.initial_situation},
        {"title": "Целевая группа", "body": output.target_group},
        {"title": "Социальная проблема", "body": output.social_problem},
        {"title": "Почему проблема значима", "body": output.significance},
        {"title": "Доказательства и источники", "body": "\n\n".join(evidence_lines)},
        {"title": "Как проект отвечает на проблему", "body": output.project_response},
        {
            "title": "Что узнать собственным опросом",
            "body": "\n".join(
                [
                    f"Гипотеза: {output.survey.hypothesis}",
                    f"Вопрос: {output.survey.question}",
                    *[f"— {answer}" for answer in output.survey.answers],
                    f"Интерпретация: {output.survey.interpretation}",
                ]
            ),
        },
    ]


def _render_docx(path: Path, output: SocialResearchOutput, sources: list[VerifiedSource], inputs: dict) -> None:
    document = Document()
    _configure_document(document)
    title = document.add_heading(output.document_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    metadata = document.add_table(rows=0, cols=2)
    metadata.style = "Table Grid"
    _table_row(metadata, "Регион", str(inputs.get("region") or "Не указан"))
    _table_row(metadata, "Направление", str(inputs.get("direction") or "Не указано"))
    _table_row(metadata, "Целевая группа", output.target_group)

    for heading, body in [
        ("1. Исходная ситуация", output.initial_situation),
        ("2. Целевая группа", output.target_group),
        ("3. Формулировка социальной проблемы", output.social_problem),
        ("4. Почему проблема значима", output.significance),
    ]:
        document.add_heading(heading, level=2)
        document.add_paragraph(_clean_text(body))

    document.add_heading("5. Доказательства и источники", level=2)
    evidence = document.add_table(rows=1, cols=4)
    evidence.style = "Table Grid"
    for cell, label in zip(evidence.rows[0].cells, ["Факт", "Источник и год", "Ссылка", "Значение для проекта"]):
        cell.text = label
    source_by_id = {source.source_id: source for source in sources}
    for item in output.evidence_items:
        source = source_by_id[item.source_id]
        cells = evidence.add_row().cells
        cells[0].text = _clean_text(item.claim)
        cells[1].text = f"{source.publisher}, {source.publication_date}"
        _add_hyperlink(cells[2].paragraphs[0], source.url, "Открыть источник")
        cells[3].text = _clean_text(item.application)

    document.add_heading("6. Как проект отвечает на проблему", level=2)
    document.add_paragraph(_clean_text(output.project_response))

    document.add_heading("7. Что узнать собственным опросом", level=2)
    document.add_paragraph(f"Проверяемая гипотеза: {_clean_text(output.survey.hypothesis)}")
    document.add_paragraph(f"Готовый вопрос для публикации во ВКонтакте: {_clean_text(output.survey.question)}")
    for answer in output.survey.answers:
        document.add_paragraph(_clean_text(answer), style="List Bullet")
    document.add_paragraph(f"Что покажет результат: {_clean_text(output.survey.interpretation)}")

    document.add_heading("8. Полезные материалы", level=2)
    useful = document.add_table(rows=1, cols=4)
    useful.style = "Table Grid"
    for cell, label in zip(useful.rows[0].cells, ["Материал", "Организация", "Год", "Ссылка"]):
        cell.text = label
    for source_id in output.useful_material_ids:
        source = source_by_id[source_id]
        cells = useful.add_row().cells
        cells[0].text = source.title
        cells[1].text = source.publisher
        cells[2].text = source.publication_date[:4]
        _add_hyperlink(cells[3].paragraphs[0], source.url, "Открыть")

    _add_page_numbers(document)
    document.save(path)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    for style_name, size in [("Title", 20), ("Heading 1", 20), ("Heading 2", 14)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(30, 64, 175) if "Heading" in style_name else RGBColor(17, 24, 39)
    if "Lary metadata" not in styles:
        metadata_style = styles.add_style("Lary metadata", WD_STYLE_TYPE.PARAGRAPH)
        metadata_style.font.name = "Arial"
        metadata_style.font.size = Pt(10)


def _table_row(table, label: str, value: str) -> None:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = _clean_text(value)


def _clean_text(value: str) -> str:
    text = re.sub(r"[`#]+", "", str(value or ""))
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return re.sub(r"\s+", " ", text).strip()


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1E40AF")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, end])


def _numbers(text: str) -> set[str]:
    return {re.sub(r"\s+", "", match) for match in re.findall(r"\d[\d\s]*(?:[.,]\d+)?", text)}


KNOWN_LOCATIONS = {
    "екатеринбург",
    "москва",
    "санкт-петербург",
    "казань",
    "краснодар",
    "кемерово",
    "нижний новгород",
    "свердловская область",
    "республика татарстан",
}


def _locations(text: str) -> set[str]:
    normalized = text.casefold().replace("ё", "е")
    return {location for location in KNOWN_LOCATIONS if location.replace("ё", "е") in normalized}
