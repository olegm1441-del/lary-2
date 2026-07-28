from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from json_repair import loads as repair_json_loads
from pydantic import BaseModel, Field

from app.services import ai_router


USER_FRIENDLY_SCENARIO_ERROR = (
    "Не удалось подготовить согласованный сценарный план. Черновик сохранён. "
    "Проверьте расписание и попробуйте повторить."
)


class ScenarioPlanGenerationError(ValueError):
    pass


class ScenarioBlock(BaseModel):
    start: str = Field(pattern=r"^\d{2}:\d{2}$")
    end: str = Field(pattern=r"^\d{2}:\d{2}$")
    duration_minutes: int = Field(gt=0, le=1440)
    title: str = Field(min_length=3, max_length=160)
    content: str = Field(min_length=10, max_length=900)
    responsible: str = Field(min_length=2, max_length=240)
    location: str = Field(min_length=2, max_length=240)
    technical_requirements: str = Field(min_length=2, max_length=500)


class ScenarioDay(BaseModel):
    day_number: int = Field(gt=0, le=30)
    day_title: str = Field(min_length=3, max_length=160)
    blocks: list[ScenarioBlock] = Field(min_length=3, max_length=40)


class PreparationStep(BaseModel):
    period: str = Field(min_length=2, max_length=120)
    actions: str = Field(min_length=10, max_length=700)
    responsible: str = Field(min_length=2, max_length=240)


class LogisticsItem(BaseModel):
    item: str = Field(min_length=2, max_length=160)
    requirement: str = Field(min_length=10, max_length=600)
    responsible: str = Field(min_length=2, max_length=240)


class ScenarioPlanOutput(BaseModel):
    document_title: str = Field(min_length=8, max_length=200)
    concept: str = Field(min_length=20, max_length=1400)
    participants: str = Field(min_length=3, max_length=700)
    beneficiary_audience: str = Field(min_length=3, max_length=700)
    capacity_summary: str = Field(min_length=10, max_length=700)
    days: list[ScenarioDay] = Field(min_length=1, max_length=30)
    preparation_steps: list[PreparationStep] = Field(min_length=1, max_length=20)
    logistics: list[LogisticsItem] = Field(min_length=1, max_length=20)
    constraints_reflected: list[str] = Field(default_factory=list, max_length=20)


@dataclass(frozen=True)
class ScenarioPlanDocument:
    output_path: Path
    output: ScenarioPlanOutput
    sections: list[dict[str, str]]


def validate_scenario_plan_output(output: ScenarioPlanOutput, inputs: dict) -> ScenarioPlanOutput:
    event_title = str(inputs.get("event_title") or "")
    event_title_tokens = _semantic_tokens(event_title)
    if event_title_tokens and not event_title_tokens.issubset(_semantic_tokens(output.document_title)):
        raise ValueError("Название документа должно использовать название мероприятия пользователя.")

    expected_days = _expected_days(str(inputs.get("schedule") or ""))
    if expected_days and len(output.days) != expected_days:
        raise ValueError(f"Сценарный план должен содержать ровно {expected_days} дня.")

    expected_numbers = list(range(1, len(output.days) + 1))
    if [day.day_number for day in output.days] != expected_numbers:
        raise ValueError("Дни должны идти последовательно без пропусков.")

    for day in output.days:
        previous_end = -1
        titles = " ".join(block.title.casefold() for block in day.blocks)
        scenario_type = str(inputs.get("scenario_type") or "").casefold()
        entry_markers = ("сбор", "инструктаж", "подготов") if any(
            marker in scenario_type for marker in ("видео", "ролик", "фильм")
        ) else ("регистрац", "сбор", "встреч")
        required_markers = (
            entry_markers,
            ("перерыв",),
            ("заверш", "закрыт", "подвед", "итог", "финал", "рефлекс", "обратн"),
        )
        for alternatives in required_markers:
            marker = "/".join(alternatives)
            if not any(alternative in titles for alternative in alternatives):
                raise ValueError(f"В каждом дне нужен операционный блок: {marker}.")
        for block in day.blocks:
            start = _minutes(block.start)
            end = _minutes(block.end)
            if start < previous_end:
                raise ValueError("Блоки сценарного плана пересекаются.")
            if end <= start or end - start != block.duration_minutes:
                raise ValueError("Длительность блока не соответствует времени начала и окончания.")
            previous_end = end

    if _normalized(output.participants) == _normalized(output.beneficiary_audience):
        raise ValueError("Участники мероприятия и целевая аудитория проекта должны быть описаны отдельно.")

    text_fields = [
        output.document_title,
        output.concept,
        output.participants,
        output.beneficiary_audience,
        output.capacity_summary,
        *(block.content for day in output.days for block in day.blocks),
        *(block.responsible for day in output.days for block in day.blocks),
        *(block.technical_requirements for day in output.days for block in day.blocks),
        *(step.actions for step in output.preparation_steps),
        *(step.responsible for step in output.preparation_steps),
        *(item.requirement for item in output.logistics),
        *(item.responsible for item in output.logistics),
        *output.constraints_reflected,
    ]
    unsupported_numbers = _numbers(" ".join(text_fields)) - _numbers(" ".join(str(value) for value in inputs.values()))
    if unsupported_numbers:
        raise ValueError(f"Сценарный план содержит неподтверждённые числа: {', '.join(sorted(unsupported_numbers))}.")

    input_location = _normalized_words(str(inputs.get("location") or ""))
    input_location_tokens = _location_tokens(str(inputs.get("location") or ""))
    if input_location:
        for day in output.days:
            for block in day.blocks:
                output_location = _normalized_words(block.location)
                shared_location_tokens = input_location_tokens.intersection(_location_tokens(block.location))
                if input_location not in output_location and len(shared_location_tokens) < min(2, len(input_location_tokens)):
                    raise ValueError("Сценарный план содержит неподтверждённое место проведения.")

    schedule_end = _schedule_end(str(inputs.get("schedule") or ""))
    if schedule_end is not None:
        for day in output.days:
            if max(_minutes(block.end) for block in day.blocks) > schedule_end:
                raise ValueError("Сценарный план выходит за указанное пользователем время.")

    constraints = str(inputs.get("team_equipment_constraints") or "").strip()
    if constraints:
        reflected = " ".join(output.constraints_reflected)
        reflected += " " + " ".join(
            f"{block.responsible} {block.technical_requirements} {block.end}"
            for day in output.days
            for block in day.blocks
        )
        reflected_tokens = _tokens(reflected)
        for clause in re.split(r"[;\n]", constraints):
            clause_tokens = _tokens(clause)
            if clause_tokens and not clause_tokens.intersection(reflected_tokens):
                raise ValueError("Не все ограничения пользователя отражены в плане.")
    return output


def build_scenario_plan_document(
    inputs: dict,
    *,
    output_path: Path | None = None,
    ai_generator: Callable[[str], str] | None = None,
) -> ScenarioPlanDocument:
    ai_generator = ai_generator or (
        lambda prompt: ai_router.generate_json_with_gigachat(prompt, ScenarioPlanOutput)
    )
    prompt = _build_prompt(inputs)
    error: Exception | None = None
    for attempt in range(2):
        try:
            raw = ai_generator(prompt if attempt == 0 else _repair_prompt(prompt, error))
            output = ScenarioPlanOutput.model_validate(_parse_scenario_json(raw, inputs))
            validate_scenario_plan_output(output, inputs)
            break
        except Exception as exc:  # provider/parser details remain internal
            error = exc
    else:
        raise ScenarioPlanGenerationError(USER_FRIENDLY_SCENARIO_ERROR) from error

    path = output_path or Path("scenario-plan.docx")
    path.parent.mkdir(parents=True, exist_ok=True)
    _render_docx(path, output, inputs)
    return ScenarioPlanDocument(path, output, _result_sections(output))


def _build_prompt(inputs: dict) -> str:
    return (
        "Ты профессиональный режиссёр и грантрайтер ПФКИ. Составь реализуемый сценарный план только по данным пользователя. "
        "Не придумывай площадки, даты, численность, команду или оборудование. Участники события и благополучатели проекта "
        "должны быть описаны отдельно. Для каждого дня нужны регистрация, содержательная программа, перерыв и завершение. "
        "Для съёмочного сценария вместо регистрации используй сбор команды или инструктаж. "
        "Если в расписании указано N дней, верни ровно N объектов days. Для каждого дня дай ровно 4 блока: "
        "регистрация или сбор, основная программа, перерыв, завершение или итоги. "
        "В content используй одно короткое предложение до 200 символов. "
        "Временные блоки не пересекаются; duration_minutes равен разнице start и end. Отрази все ограничения пользователя. "
        "Не оставляй обязательные строки пустыми: если техника не нужна, напиши «Не требуется»; "
        "для каждого блока укажи ответственного из команды пользователя или нейтральную проектную роль без нового имени. "
        "Верни только JSON без markdown по схеме: "
        '{"document_title":"...","concept":"...","participants":"...","beneficiary_audience":"...",'
        '"capacity_summary":"...","days":[{"day_number":1,"day_title":"...",'
        '"blocks":[{"start":"10:00","end":"10:30","duration_minutes":30,"title":"...",'
        '"content":"...","responsible":"...","location":"...","technical_requirements":"..."}]}],'
        '"preparation_steps":[{"period":"...","actions":"...","responsible":"..."}],'
        '"logistics":[{"item":"...","requirement":"...","responsible":"..."}],'
        '"constraints_reflected":["..."]}. '
        f"Входные данные: {json.dumps(inputs, ensure_ascii=False, sort_keys=True)}"
    )


def _repair_prompt(prompt: str, error: Exception | None) -> str:
    reason = str(error or "неизвестная ошибка")[:500]
    return (
        f"{prompt}\nПредыдущий ответ не прошёл проверку: {reason}. "
        "Исправь только структуру и противоречия. Верни один валидный JSON без пояснений."
    )


def _extract_json(raw: str) -> str:
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", str(raw).strip(), flags=re.IGNORECASE)
    start, end = cleaned.find("{"), cleaned.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("AI не вернул JSON.")
    return cleaned[start : end + 1]


def _parse_scenario_json(raw: str, inputs: dict | None = None) -> dict:
    parsed = repair_json_loads(_extract_json(raw))
    if not isinstance(parsed, dict):
        raise ValueError("AI не вернул JSON-объект.")
    days = parsed.get("days")
    if isinstance(days, list) and days and isinstance(days[-1], dict):
        tail = days[-1]
        root_section_keys = {"preparation_steps", "logistics", "constraints_reflected"}
        if (
            "day_number" not in tail
            and set(tail).issubset(root_section_keys)
            and set(tail).intersection({"preparation_steps", "logistics"})
        ):
            for key in root_section_keys:
                if key in tail and key not in parsed:
                    parsed[key] = tail[key]
            days.pop()
    if isinstance(days, list):
        for day in days:
            if not isinstance(day, dict) or not isinstance(day.get("blocks"), list):
                continue
            for block in day["blocks"]:
                if not isinstance(block, dict):
                    continue
                start, end = block.get("start"), block.get("end")
                if not isinstance(start, str) or not isinstance(end, str):
                    continue
                try:
                    start_minutes, end_minutes = _minutes(start), _minutes(end)
                except (TypeError, ValueError):
                    continue
                if end_minutes > start_minutes:
                    block["duration_minutes"] = end_minutes - start_minutes
    if inputs is not None:
        schedule = str(inputs.get("schedule") or "")
        schedule_window = _schedule_window(schedule)
        if schedule_window is not None and isinstance(days, list):
            for day in days:
                if not isinstance(day, dict) or not isinstance(day.get("blocks"), list):
                    continue
                if len(day["blocks"]) == 4:
                    _distribute_four_blocks(day["blocks"], *schedule_window)
        schedule_end = _schedule_end(schedule)
        if schedule_end is not None and isinstance(days, list):
            for day in days:
                if not isinstance(day, dict) or not isinstance(day.get("blocks"), list):
                    continue
                for block in day["blocks"]:
                    if not isinstance(block, dict):
                        continue
                    try:
                        start_minutes = _minutes(str(block.get("start") or ""))
                        end_minutes = _minutes(str(block.get("end") or ""))
                    except (TypeError, ValueError):
                        continue
                    if end_minutes > schedule_end and start_minutes < schedule_end:
                        block["end"] = f"{schedule_end // 60:02d}:{schedule_end % 60:02d}"
                        block["duration_minutes"] = schedule_end - start_minutes
        location = str(inputs.get("location") or "").strip()
        if isinstance(days, list):
            for day in days:
                if not isinstance(day, dict) or not isinstance(day.get("blocks"), list):
                    continue
                for block in day["blocks"]:
                    if not isinstance(block, dict):
                        continue
                    if location:
                        block["location"] = location
                    elif not str(block.get("location") or "").strip():
                        block["location"] = "Площадка мероприятия"
                    if not str(block.get("responsible") or "").strip():
                        block["responsible"] = "Команда проекта"
                    if not str(block.get("technical_requirements") or "").strip():
                        block["technical_requirements"] = "Не требуется"
        for field in ("participants", "beneficiary_audience"):
            user_value = str(inputs.get(field) or "").strip()
            if user_value:
                parsed[field] = user_value
        preparation = str(inputs.get("preparation") or "").strip()
        if not parsed.get("preparation_steps"):
            parsed["preparation_steps"] = [
                {
                    "period": "До начала мероприятия",
                    "actions": preparation or "Подготовить площадку и проверить готовность программы.",
                    "responsible": "Команда проекта",
                }
            ]
        else:
            for step in parsed["preparation_steps"]:
                if not isinstance(step, dict):
                    continue
                if len(str(step.get("period") or "").strip()) < 2:
                    step["period"] = "До начала мероприятия"
                if len(str(step.get("actions") or "").strip()) < 10:
                    step["actions"] = preparation or "Подготовить площадку и проверить готовность программы."
                if len(str(step.get("responsible") or "").strip()) < 2:
                    step["responsible"] = "Команда проекта"
        constraints = str(inputs.get("team_equipment_constraints") or "").strip()
        safe_logistics_requirement = f"Организовать работу на площадке «{location}»"
        if constraints:
            safe_logistics_requirement += f" с учётом условий: {constraints}"
        safe_logistics_requirement = safe_logistics_requirement.rstrip(".") + "."
        if not parsed.get("logistics"):
            parsed["logistics"] = [
                {
                    "item": "Организация площадки",
                    "requirement": safe_logistics_requirement,
                    "responsible": "Команда проекта",
                }
            ]
        else:
            for item in parsed["logistics"]:
                if not isinstance(item, dict):
                    continue
                if len(str(item.get("item") or "").strip()) < 2:
                    item["item"] = "Организация площадки"
                if len(str(item.get("requirement") or "").strip()) < 10:
                    item["requirement"] = safe_logistics_requirement
                if len(str(item.get("responsible") or "").strip()) < 2:
                    item["responsible"] = "Команда проекта"
        if constraints and not parsed.get("constraints_reflected"):
            parsed["constraints_reflected"] = [
                clause.strip()
                for clause in re.split(r"[;\n]", constraints)
                if clause.strip()
            ]
    return parsed


def _expected_days(schedule: str) -> int | None:
    match = re.search(r"\b(\d{1,2})\s*(?:дн(?:я|ей|ь)|дня)\b", schedule.casefold())
    return int(match.group(1)) if match else None


def _schedule_end(schedule: str) -> int | None:
    matches = re.findall(r"(?:до|–|-)\s*(\d{1,2}:\d{2})", schedule)
    return _minutes(matches[-1]) if matches else None


def _schedule_window(schedule: str) -> tuple[int, int] | None:
    matches = re.findall(r"\d{1,2}:\d{2}", schedule)
    if len(matches) < 2:
        return None
    start, end = _minutes(matches[0]), _minutes(matches[-1])
    return (start, end) if end - start >= 120 else None


def _distribute_four_blocks(blocks: list[dict], start: int, end: int) -> None:
    entry_end = min(start + 30, end - 90)
    midpoint = start + (end - start) // 2
    break_start = max(entry_end + 30, (midpoint // 30) * 30)
    break_start = min(break_start, end - 60)
    break_end = min(break_start + 30, end - 30)
    boundaries = [start, entry_end, break_start, break_end, end]
    for index, block in enumerate(blocks):
        block_start, block_end = boundaries[index], boundaries[index + 1]
        block["start"] = f"{block_start // 60:02d}:{block_start % 60:02d}"
        block["end"] = f"{block_end // 60:02d}:{block_end % 60:02d}"
        block["duration_minutes"] = block_end - block_start


def _minutes(value: str) -> int:
    hours, minutes = (int(part) for part in value.split(":"))
    if hours > 23 or minutes > 59:
        raise ValueError("Некорректное время.")
    return hours * 60 + minutes


def _normalized(value: str) -> str:
    return re.sub(r"\W+", "", value.casefold())


def _normalized_words(value: str) -> str:
    return " ".join(re.findall(r"[а-яёa-z0-9]+", value.casefold().replace("ё", "е")))


def _location_tokens(value: str) -> set[str]:
    return {word[:4] for word in _normalized_words(value).split() if len(word) >= 4}


def _semantic_tokens(value: str) -> set[str]:
    return {word[:4] for word in _normalized_words(value).split() if len(word) >= 4}


def _numbers(value: str) -> set[str]:
    return {re.sub(r"\s+", "", match) for match in re.findall(r"\d[\d\s]*(?:[.,]\d+)?", value)}


def _tokens(value: str) -> set[str]:
    return {item.casefold().replace("ё", "е") for item in re.findall(r"[\w:]+", value) if len(item) >= 2}


def _result_sections(output: ScenarioPlanOutput) -> list[dict[str, str]]:
    sections = [
        {"title": "Концепция", "body": output.concept},
        {"title": "Участники мероприятия", "body": output.participants},
        {"title": "Целевая аудитория проекта", "body": output.beneficiary_audience},
        {"title": "Вместимость и состав", "body": output.capacity_summary},
    ]
    for day in output.days:
        body = "\n".join(
            f"{block.start}–{block.end}. {block.title}. {block.content} Ответственные: {block.responsible}."
            for block in day.blocks
        )
        sections.append({"title": day.day_title, "body": body})
    return sections


def _render_docx(path: Path, output: ScenarioPlanOutput, inputs: dict) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    _configure_styles(document)

    document.add_heading(output.document_title, level=0)
    metadata = document.add_table(rows=0, cols=2)
    metadata.style = "Table Grid"
    _table_row(metadata, "Формат", str(inputs.get("scenario_type") or "—"))
    _table_row(metadata, "Место проведения", str(inputs.get("location") or "—"))
    _table_row(metadata, "Расписание", str(inputs.get("schedule") or "—"))
    _table_row(metadata, "Участники мероприятия", output.participants)
    _table_row(metadata, "Целевая аудитория проекта", output.beneficiary_audience)

    document.add_heading("Концепция", level=1)
    document.add_paragraph(output.concept)
    document.add_paragraph(output.capacity_summary)

    for day in output.days:
        document.add_heading(day.day_title, level=1)
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.autofit = False
        widths = [Cm(2.3), Cm(3.4), Cm(7.2), Cm(4.3), Cm(3.8), Cm(4.7)]
        headers = ["Время", "Блок", "Содержание", "Ответственные", "Место", "Техника"]
        for cell, label in zip(table.rows[0].cells, headers):
            cell.text = label
        _set_row_widths(table.rows[0], widths)
        _repeat_table_header(table.rows[0])
        for block in day.blocks:
            cells = table.add_row().cells
            values = [
                f"{block.start}–{block.end}",
                block.title,
                block.content,
                block.responsible,
                block.location,
                block.technical_requirements,
            ]
            for cell, value in zip(cells, values):
                cell.text = value
            _set_row_widths(table.rows[-1], widths)

    document.add_heading("Подготовка", level=1)
    preparation = document.add_table(rows=1, cols=3)
    preparation.style = "Table Grid"
    for cell, label in zip(preparation.rows[0].cells, ["Период", "Действия", "Ответственные"]):
        cell.text = label
    for step in output.preparation_steps:
        cells = preparation.add_row().cells
        for cell, value in zip(cells, [step.period, step.actions, step.responsible]):
            cell.text = value

    document.add_heading("Логистика", level=1)
    logistics = document.add_table(rows=1, cols=3)
    logistics.style = "Table Grid"
    for cell, label in zip(logistics.rows[0].cells, ["Задача", "Требование", "Ответственные"]):
        cell.text = label
    for item in output.logistics:
        cells = logistics.add_row().cells
        for cell, value in zip(cells, [item.item, item.requirement, item.responsible]):
            cell.text = value

    _style_table_text(document)
    _add_page_numbers(document)
    document.save(path)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [
        ("Title", 20, RGBColor(10, 25, 56)),
        ("Heading 1", 15, RGBColor(17, 65, 125)),
    ]:
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
    if "Table Text" not in document.styles:
        table_style = document.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
        table_style.font.name = "Arial"
        table_style.font.size = Pt(9)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = document.styles["Table Text"]


def _table_row(table, label: str, value: str) -> None:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = value


def _set_row_widths(row, widths: list[Cm]) -> None:
    for cell, width in zip(row.cells, widths):
        cell.width = width


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _style_table_text(document: Document) -> None:
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)
                        if row_index == 0:
                            run.font.bold = True


def _add_page_numbers(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Страница ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)
