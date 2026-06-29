import json
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document

from app.services.support_letter import (
    SupportLetterGenerationError,
    SupportLetterValidationError,
    build_support_letter_document,
    normalize_cofinance,
    normalize_partner_name,
    normalize_project_title,
    normalize_support_types,
    sanitize_filename,
)


VALID_PAYLOAD = {
    "contest": "pfki",
    "project_title": "«Фестиваль \"Теплый дом\"»",
    "partner_name": "ООО \"Лютики\"",
    "partner_intro_block": "крупнейший региональный поставщик футболок.",
    "value_keywords": "дети 10–17 лет; Екатеринбург; творческое самовыражение; уверенность; общение со сверстниками; семейное участие",
    "support_types": ["Информационная", "Материальная"],
    "support_details": "разместит 5 публикаций в социальных сетях и предоставит 100 футболок для победителей и участников финального мероприятия",
    "cofinance_block": "600000",
    "signatory": "Генеральный директор ООО \"Лютики\" Иванов Иван Иванович",
}


def docx_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


class SupportLetterServiceTest(unittest.TestCase):
    def test_normalization_examples_match_pfki_template_rules(self):
        self.assertEqual(normalize_partner_name('ООО "Лютики"'), "ООО «Лютики»")
        self.assertEqual(normalize_partner_name("ООО «Лютики»"), "ООО «Лютики»")
        self.assertEqual(normalize_partner_name('МБУК "Городской дом культуры"'), "МБУК «Городской дом культуры»")
        self.assertEqual(normalize_partner_name('ООО "кринжульки"'), "ООО «Кринжульки»")
        self.assertEqual(normalize_partner_name("ООО «кринжульки»"), "ООО «Кринжульки»")
        self.assertEqual(normalize_partner_name('АНО "теплый дом"'), "АНО «Теплый дом»")

        self.assertEqual(normalize_project_title('"Детство. Лагерь.Like"'), "Детство. Лагерь.Like")
        self.assertEqual(normalize_project_title("«Детство. Лагерь.Like»"), "Детство. Лагерь.Like")
        self.assertEqual(normalize_project_title('Семейный фестиваль "Теплый дом"'), "Семейный фестиваль „Теплый дом“")
        self.assertEqual(normalize_project_title("«Семейный фестиваль «Теплый дом»»"), "Семейный фестиваль „Теплый дом“")
        self.assertEqual(normalize_project_title('"домик в деревне"'), "Домик в деревне")
        self.assertEqual(normalize_project_title("«домик в деревне»"), "Домик в деревне")
        self.assertEqual(normalize_project_title('Фестиваль "теплый дом"'), "Фестиваль „Теплый дом“")
        self.assertEqual(normalize_project_title("«Фестиваль «теплый дом»»"), "Фестиваль „Теплый дом“")

        cofinance = normalize_cofinance("600000")
        self.assertEqual(cofinance.raw_digits, "600000")
        self.assertEqual(cofinance.formatted, "600 000")
        self.assertEqual(normalize_cofinance("600 000").formatted, "600 000")

        with self.assertRaises(SupportLetterValidationError):
            normalize_cofinance("600000 рублей")

        self.assertEqual(sanitize_filename("ООО «Лютики»"), "Письмо поддержки_ПФКИ_ООО «Лютики».docx")
        self.assertEqual(sanitize_filename('ООО "Лютики" / тест'), "Письмо поддержки_ПФКИ_ООО Лютики тест.docx")
        self.assertEqual(sanitize_filename("____"), "Письмо поддержки_ПФКИ_партнер.docx")

        self.assertEqual(normalize_support_types(["Информационная", "Материальная"]), ["Информационная", "Материальная"])
        with self.assertRaises(SupportLetterValidationError):
            normalize_support_types(["Подарки / призы"])

    def test_support_letter_docx_uses_template_and_two_short_ai_json_blocks(self):
        ai_value = {
            "ai_value_block": "Видим необходимость проекта в следующем:\n1. Проект помогает детям безопасно проявлять себя творчески.\n2. Проект усиливает семейное участие и общение со сверстниками.\n3. Проект создает понятный культурный маршрут для участников.\nВидим особенным этот проект для нашей территории Екатеринбурга.",
        }
        ai_support = {
            "ai_support_block": "Информационная поддержка: разместим 5 публикаций в социальных сетях проекта. Подготовим короткий анонс и напоминание для участников.\nМатериальная поддержка: предоставим 100 футболок для победителей и участников финала. Передадим их команде проекта до события.",
        }

        with patch("app.services.support_letter.generate_with_gigachat", side_effect=[json.dumps(ai_value, ensure_ascii=False), json.dumps(ai_support, ensure_ascii=False)]) as ai:
            result = build_support_letter_document(VALID_PAYLOAD)

        self.assertEqual(ai.call_count, 2)
        prompts = "\n\n".join(call.args[0] for call in ai.call_args_list)
        self.assertIn("Составь блок {{AI_VALUE_BLOCK}}", prompts)
        self.assertIn("Составь блок {{AI_SUPPORT_BLOCK}}", prompts)
        self.assertIn("Выбранные виды поддержки: Информационная; Материальная", prompts)
        self.assertNotIn(VALID_PAYLOAD["signatory"], prompts)
        self.assertIn("Организация-партнер НЕ является заявителем", prompts)

        text = docx_text(result.docx_bytes)
        self.assertNotIn("{{", text)
        self.assertIn("Мы – ООО «Лютики» – крупнейший региональный поставщик футболок. Выражаем поддержку проекта «Фестиваль „Теплый дом“».", text)
        self.assertIn("Со своей стороны, готовы оказать проекту следующую поддержку:", text)
        self.assertIn("Оцениваем наш вклад в 600 000 рублей", text)
        self.assertNotIn("600 000 рублей рублей", text)
        self.assertIn("Генеральный директор ООО «Лютики» Иванов Иван Иванович", text)
        self.assertEqual(result.filename, "Письмо поддержки_ПФКИ_ООО «Лютики».docx")

    def test_empty_deterministic_fields_render_placeholders_and_filename_fallback(self):
        payload = {
            "contest": "pfki",
            "project_title": "",
            "partner_name": "",
            "partner_intro_block": "",
            "value_keywords": "дети 10–17 лет; Екатеринбург; творческое самовыражение; уверенность; семейное участие",
            "support_types": ["Информационная"],
            "support_details": "размещение публикаций о проекте на информационных площадках партнера",
            "cofinance_block": "",
            "signatory": "",
        }
        ai_value = {
            "ai_value_block": "Видим необходимость проекта в следующем:\n1. Проект помогает детям проявлять себя творчески.\n2. Проект поддерживает уверенность участников.\n3. Проект вовлекает семьи в культурную активность.\nВидим особенным этот проект для нашей территории.",
        }
        ai_support = {
            "ai_support_block": "Информационная поддержка: готовы разместить публикации о проекте на информационных площадках. Это поможет донести информацию до потенциальных участников.",
        }

        with patch("app.services.support_letter.generate_with_gigachat", side_effect=[json.dumps(ai_value, ensure_ascii=False), json.dumps(ai_support, ensure_ascii=False)]) as ai:
            result = build_support_letter_document(payload)

        prompts = "\n\n".join(call.args[0] for call in ai.call_args_list)
        self.assertIn("Название проекта: Название проекта не указано", prompts)
        self.assertIn("Организация-партнер: Организация-партнер не указана", prompts)
        text = docx_text(result.docx_bytes)
        self.assertNotIn("{{", text)
        self.assertIn("____", text)
        self.assertIn("Оцениваем наш вклад в ____ рублей", text)
        self.assertEqual(result.filename, "Письмо поддержки_ПФКИ_партнер.docx")

    def test_forbidden_partner_as_implementer_wording_triggers_retry(self):
        bad_value = {
            "ai_value_block": "Видим необходимость проекта, реализуемого ООО «Кринжульки», для пожилых жителей территории.",
        }
        good_value = {
            "ai_value_block": "Видим необходимость проекта в следующем:\n1. Проект поддерживает пожилых жителей территории.\n2. Проект помогает людям старшего поколения сохранить активность.\n3. Проект усиливает социальные связи участников.\nВидим особенным этот проект для нашей территории.",
        }
        good_support = {
            "ai_support_block": "Информационная поддержка: готовы разместить публикации о проекте на своих площадках. Это поможет привлечь внимание участников.",
        }
        payload = {
            **VALID_PAYLOAD,
            "project_title": "«домик в деревне»",
            "partner_name": 'ООО "кринжульки"',
            "support_types": ["Информационная"],
            "support_details": "разместит публикации о проекте на своих информационных площадках до конца 2027 года",
        }

        with patch(
            "app.services.support_letter.generate_with_gigachat",
            side_effect=[json.dumps(bad_value, ensure_ascii=False), json.dumps(good_value, ensure_ascii=False), json.dumps(good_support, ensure_ascii=False)],
        ) as ai:
            result = build_support_letter_document(payload)

        self.assertEqual(ai.call_count, 3)
        retry_prompt = ai.call_args_list[1].args[0]
        self.assertIn("Партнер только оказывает поддержку", retry_prompt)
        self.assertNotIn("реализуемого ООО «Кринжульки»", result.ai_value_block)
        self.assertEqual(result.normalized.project_title, "Домик в деревне")
        self.assertEqual(result.normalized.partner_name, "ООО «Кринжульки»")
        self.assertEqual(result.filename, "Письмо поддержки_ПФКИ_ООО «Кринжульки».docx")

    def test_unselected_support_type_in_ai_block_triggers_soft_error_after_retry(self):
        bad_support = {
            "ai_support_block": "Финансовая поддержка: готовы оплатить расходы проекта. Информационная поддержка: готовы разместить анонс проекта.",
        }
        good_value = {
            "ai_value_block": "Видим необходимость проекта в следующем:\n1. Проект помогает детям проявлять себя творчески.\n2. Проект поддерживает уверенность участников.\n3. Проект вовлекает семьи в культурную активность.\nВидим особенным этот проект для нашей территории.",
        }
        payload = {**VALID_PAYLOAD, "support_types": ["Информационная"]}

        with patch(
            "app.services.support_letter.generate_with_gigachat",
            side_effect=[json.dumps(good_value, ensure_ascii=False), json.dumps(bad_support, ensure_ascii=False), json.dumps(bad_support, ensure_ascii=False)],
        ):
            with self.assertRaises(SupportLetterGenerationError):
                build_support_letter_document(payload)


if __name__ == "__main__":
    unittest.main()
