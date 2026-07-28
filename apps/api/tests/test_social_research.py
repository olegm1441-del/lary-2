import json
import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

os.environ.setdefault("APP_ENV", "test")

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services.account_store import clear_account_store_for_tests
from app.services.social_research import (
    SocialResearchGenerationError,
    SocialResearchOutput,
    VerifiedSource,
    USER_FRIENDLY_SOURCE_ERROR,
    build_social_research_document,
    validate_social_research_output,
    validate_verified_sources,
)


INPUTS = {
    "region": "Свердловская область",
    "direction": "Театр",
    "target_group": "Подростки и молодёжь 12–22 лет в городах-миллионниках",
    "problem": "Молодёжь плохо разбирается в современных форматах народного театра.",
    "project_response": "Проект знакомит молодёжь с возможностями театрального досуга и современными форматами народного театра.",
    "constraints": "Можно использовать только офлайн-инструменты.",
}


def source(source_id: str, *, date: str = "2024-02-16", url: str | None = None) -> VerifiedSource:
    return VerifiedSource(
        source_id=source_id,
        title=f"Материал {source_id}",
        publisher="Официальная организация",
        publication_date=date,
        url=url or f"https://example.org/{source_id}",
        source_type="official" if source_id == "source-1" else "research",
        territory="Российская Федерация",
        claim="Проверенный вывод без дополнительных чисел.",
        evidence="Публикация содержит сведения по теме проекта.",
        verified_at="2026-07-28T12:00:00+00:00",
    )


def output_payload() -> dict:
    return {
        "document_title": "Актуальность проекта о современных форматах народного театра",
        "initial_situation": "Целевая группа недостаточно знакома с доступными театральными форматами.",
        "target_group": INPUTS["target_group"],
        "social_problem": INPUTS["problem"],
        "significance": "Недостаток информации ограничивает осознанный выбор культурного досуга.",
        "project_response": INPUTS["project_response"],
        "evidence_items": [
            {
                "claim": "Проверенный источник описывает интерес аудитории к театральным форматам.",
                "source_id": "source-1",
                "application": "Использовать при обосновании актуальности проекта.",
            },
            {
                "claim": "Исследовательский материал помогает описать особенности культурного участия.",
                "source_id": "source-2",
                "application": "Использовать при описании целевой группы.",
            },
        ],
        "survey": {
            "hypothesis": "Целевая группа недостаточно знает о доступных театральных событиях.",
            "question": "Насколько хорошо вы знаете, какие театральные события проходят в вашем городе?",
            "answers": [
                "Хорошо знаю и регулярно слежу",
                "Знаю несколько событий",
                "Иногда вижу информацию случайно",
                "Почти ничего не знаю",
                "Не интересуюсь этой темой",
            ],
            "interpretation": "Ответы помогут оценить дефицит информирования, но не являются уже проведённым исследованием.",
        },
        "useful_material_ids": ["source-1", "source-2"],
        "manual_checks": ["Уточнить конкретные города, если территория проекта уже определена."],
    }


class SocialResearchContractTest(unittest.TestCase):
    def setUp(self):
        clear_account_store_for_tests()

    def test_verified_source_rejects_publication_before_2024(self):
        with self.assertRaises(ValueError):
            validate_verified_sources([source("source-1", date="2023-12-31")])

    def test_verified_source_rejects_non_http_url(self):
        with self.assertRaises(ValueError):
            validate_verified_sources([source("source-1", url="file:///tmp/source.pdf")])

    def test_verified_source_rejects_duplicate_url(self):
        duplicate = "https://example.org/material"
        with self.assertRaises(ValueError):
            validate_verified_sources([
                source("source-1", url=duplicate),
                source("source-2", url=duplicate),
            ])

    def test_ai_cannot_introduce_unknown_source_id(self):
        payload = output_payload()
        payload["evidence_items"][0]["source_id"] = "source-unknown"
        with self.assertRaises(ValueError):
            validate_social_research_output(SocialResearchOutput.model_validate(payload), [source("source-1"), source("source-2")], INPUTS)

    def test_unsupported_number_is_rejected(self):
        payload = output_payload()
        payload["significance"] += " Показатель составляет 77 процентов."
        with self.assertRaises(ValueError):
            validate_social_research_output(SocialResearchOutput.model_validate(payload), [source("source-1"), source("source-2")], INPUTS)

    def test_evidence_number_must_belong_to_its_own_source(self):
        sources = [
            source("source-1").model_copy(update={"claim": "Официальный показатель составляет 55 процентов."}),
            source("source-2"),
        ]
        payload = output_payload()
        payload["evidence_items"][1]["claim"] = "Исследовательский показатель составляет 55 процентов."
        with self.assertRaisesRegex(ValueError, "указанного источника"):
            validate_social_research_output(SocialResearchOutput.model_validate(payload), sources, INPUTS)

    def test_invented_location_is_rejected(self):
        payload = output_payload()
        payload["initial_situation"] += " Особенно заметно это в Екатеринбурге."
        with self.assertRaises(ValueError):
            validate_social_research_output(SocialResearchOutput.model_validate(payload), [source("source-1"), source("source-2")], INPUTS)

    def test_survey_requires_four_to_six_distinct_answers(self):
        payload = output_payload()
        payload["survey"]["answers"] = ["Да", "Нет", "Да"]
        with self.assertRaises(ValueError):
            SocialResearchOutput.model_validate(payload)

    def test_no_verified_source_returns_controlled_failure(self):
        with self.assertRaisesRegex(SocialResearchGenerationError, "Не удалось получить проверяемые источники"):
            build_social_research_document(INPUTS, source_fetcher=lambda _: [], ai_generator=lambda _: "{}")

    def test_document_has_verified_hyperlinks_and_no_technical_dump(self):
        sources = [source("source-1"), source("source-2")]
        payload = output_payload()
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "research.docx"
            document = build_social_research_document(
                INPUTS,
                output_path=path,
                source_fetcher=lambda _: sources,
                ai_generator=lambda _: json.dumps(payload, ensure_ascii=False),
            )

            self.assertEqual(document.output.document_title, payload["document_title"])
            self.assertTrue(path.exists())
            with ZipFile(path) as archive:
                self.assertIn("word/document.xml", archive.namelist())
                relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
                self.assertIn("https://example.org/source-1", relationships)
                self.assertIn("https://example.org/source-2", relationships)

            parsed = Document(path)
            full_text = "\n".join(p.text for p in parsed.paragraphs)
            full_text += "\n" + "\n".join(cell.text for table in parsed.tables for row in table.rows for cell in row.cells)
            for banned in ["Ищите данные", "backend", "schema", "```", "##", "{{"]:
                self.assertNotIn(banned, full_text)
            self.assertIn("Что узнать собственным опросом", full_text)
            self.assertGreaterEqual(len(parsed.tables), 2)

    def test_invalid_ai_json_gets_one_repair_attempt(self):
        sources = [source("source-1"), source("source-2")]
        responses = iter(["```json\n{broken}\n```", json.dumps(output_payload(), ensure_ascii=False)])
        with tempfile.TemporaryDirectory() as directory:
            document = build_social_research_document(
                INPUTS,
                output_path=Path(directory) / "research.docx",
                source_fetcher=lambda _: sources,
                ai_generator=lambda _: next(responses),
            )
        self.assertEqual(document.output.document_title, output_payload()["document_title"])

    def test_failed_generation_does_not_spend_free_attempt(self):
        client = TestClient(app)
        before = client.get("/api/usage").json()["modules"]["social-research"]["free_attempt_available"]
        with patch(
            "app.services.module_engine.build_social_research_document",
            side_effect=SocialResearchGenerationError(USER_FRIENDLY_SOURCE_ERROR),
        ):
            response = client.post(
                "/api/module-runs",
                json={"module_slug": "social-research", "contest_slug": "pfki", "inputs": INPUTS},
            )
        after = client.get("/api/usage").json()["modules"]["social-research"]["free_attempt_available"]
        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"]["message"], USER_FRIENDLY_SOURCE_ERROR)
        self.assertTrue(before)
        self.assertTrue(after)


if __name__ == "__main__":
    unittest.main()
