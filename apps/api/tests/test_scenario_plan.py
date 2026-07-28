import json
import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

os.environ.setdefault("APP_ENV", "test")

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services.account_store import clear_account_store_for_tests
from app.services.module_inputs import normalize_inputs
from app.services.scenario_plan import (
    ScenarioPlanGenerationError,
    ScenarioPlanOutput,
    USER_FRIENDLY_SCENARIO_ERROR,
    build_scenario_plan_document,
    validate_scenario_plan_output,
)


INPUTS = {
    "scenario_type": "Фестиваль",
    "event_title": "Фестиваль «Городские истории»",
    "event_idea": "Трёхдневный фестиваль с мастерскими, показами и обсуждениями.",
    "location": "Дом культуры в Казани",
    "participants": "20 артистов, 6 волонтёров, до 150 зрителей одновременно",
    "beneficiary_audience": "Подростки 14–17 лет и их родители",
    "schedule": "3 дня; ежедневно с 10:00 до 18:00",
    "preparation": "14 дней на набор участников, репетиции и монтаж площадки",
    "team_equipment_constraints": "2 ведущих, звукорежиссёр, проектор; окончание до 18:00",
}


def output_payload() -> dict:
    days = []
    for day_number in range(1, 4):
        days.append(
            {
                "day_number": day_number,
                "day_title": f"День {day_number}",
                "blocks": [
                    {
                        "start": "10:00",
                        "end": "10:30",
                        "duration_minutes": 30,
                        "title": "Регистрация участников",
                        "content": "Встреча участников и проверка списков.",
                        "responsible": "Координатор и волонтёры",
                        "location": "Фойе Дома культуры в Казани",
                        "technical_requirements": "Стол регистрации и списки участников",
                    },
                    {
                        "start": "10:30",
                        "end": "13:00",
                        "duration_minutes": 150,
                        "title": "Основная программа",
                        "content": "Мастерские, показы и обсуждения с подростками и родителями.",
                        "responsible": "Два ведущих и артисты",
                        "location": "Зал Дома культуры в Казани",
                        "technical_requirements": "Проектор и работа звукорежиссёра",
                    },
                    {
                        "start": "13:00",
                        "end": "13:30",
                        "duration_minutes": 30,
                        "title": "Перерыв",
                        "content": "Организованный перерыв и проветривание зала.",
                        "responsible": "Координатор",
                        "location": "Дом культуры в Казани",
                        "technical_requirements": "Навигация для участников",
                    },
                    {
                        "start": "13:30",
                        "end": "17:30",
                        "duration_minutes": 240,
                        "title": "Продолжение программы",
                        "content": "Практические занятия и итоговые показы.",
                        "responsible": "Ведущие, артисты и звукорежиссёр",
                        "location": "Зал Дома культуры в Казани",
                        "technical_requirements": "Проектор и звуковое оборудование",
                    },
                    {
                        "start": "17:30",
                        "end": "18:00",
                        "duration_minutes": 30,
                        "title": "Завершение дня",
                        "content": "Подведение итогов и организованный выход зрителей.",
                        "responsible": "Координатор и волонтёры",
                        "location": "Зал Дома культуры в Казани",
                        "technical_requirements": "Микрофон",
                    },
                ],
            }
        )
    return {
        "document_title": "Сценарный план фестиваля «Городские истории»",
        "concept": "Фестиваль объединяет мастерские, показы и обсуждения в последовательную трёхдневную программу.",
        "participants": INPUTS["participants"],
        "beneficiary_audience": INPUTS["beneficiary_audience"],
        "capacity_summary": "Одновременно площадка принимает до 150 зрителей; команда включает артистов, волонтёров, ведущих и звукорежиссёра.",
        "days": days,
        "preparation_steps": [
            {
                "period": "За 14 дней",
                "actions": "Набор участников, репетиции и согласование программы.",
                "responsible": "Координатор и ведущие",
            },
            {
                "period": "Накануне",
                "actions": "Монтаж проектора и проверка звукового оборудования.",
                "responsible": "Звукорежиссёр и техническая команда",
            },
        ],
        "logistics": [
            {
                "item": "Потоки участников",
                "requirement": "Разделить вход, регистрацию и выход зрителей.",
                "responsible": "Волонтёры",
            }
        ],
        "constraints_reflected": [
            "В программе задействованы два ведущих и звукорежиссёр.",
            "Проектор предусмотрен в технических требованиях.",
            "Все дни завершаются до 18:00.",
        ],
    }


class ScenarioPlanContractTest(unittest.TestCase):
    def setUp(self):
        clear_account_store_for_tests()

    def test_requires_exact_number_of_days_from_schedule(self):
        payload = output_payload()
        payload["days"].pop()
        with self.assertRaisesRegex(ValueError, "3"):
            validate_scenario_plan_output(ScenarioPlanOutput.model_validate(payload), INPUTS)

    def test_document_title_uses_explicit_event_title(self):
        payload = output_payload()
        payload["document_title"] = "Сценарный план другого мероприятия"
        with self.assertRaisesRegex(ValueError, "название мероприятия"):
            validate_scenario_plan_output(ScenarioPlanOutput.model_validate(payload), INPUTS)

    def test_rejects_overlapping_blocks(self):
        payload = output_payload()
        payload["days"][0]["blocks"][1]["start"] = "10:20"
        with self.assertRaisesRegex(ValueError, "пересека"):
            validate_scenario_plan_output(ScenarioPlanOutput.model_validate(payload), INPUTS)

    def test_rejects_incorrect_duration(self):
        payload = output_payload()
        payload["days"][0]["blocks"][0]["duration_minutes"] = 10
        with self.assertRaisesRegex(ValueError, "Длитель"):
            validate_scenario_plan_output(ScenarioPlanOutput.model_validate(payload), INPUTS)

    def test_requires_operational_blocks(self):
        payload = output_payload()
        payload["days"][0]["blocks"] = [
            block for block in payload["days"][0]["blocks"] if "Регистрация" not in block["title"]
        ]
        with self.assertRaisesRegex(ValueError, "регистрац"):
            validate_scenario_plan_output(ScenarioPlanOutput.model_validate(payload), INPUTS)

    def test_participants_and_beneficiaries_stay_separate(self):
        payload = output_payload()
        payload["beneficiary_audience"] = payload["participants"]
        with self.assertRaisesRegex(ValueError, "аудитор"):
            validate_scenario_plan_output(ScenarioPlanOutput.model_validate(payload), INPUTS)

    def test_rejects_invented_capacity(self):
        payload = output_payload()
        payload["capacity_summary"] = "Площадка одновременно принимает 999 зрителей и проектную команду."
        with self.assertRaisesRegex(ValueError, "неподтверждённые числа"):
            validate_scenario_plan_output(ScenarioPlanOutput.model_validate(payload), INPUTS)

    def test_invalid_ai_output_gets_one_repair_attempt(self):
        responses = iter(["```json\n{broken}\n```", json.dumps(output_payload(), ensure_ascii=False)])
        calls = []

        def generate(prompt: str) -> str:
            calls.append(prompt)
            return next(responses)

        with tempfile.TemporaryDirectory() as directory:
            result = build_scenario_plan_document(
                INPUTS,
                output_path=Path(directory) / "scenario.docx",
                ai_generator=generate,
            )
        self.assertEqual(result.output.document_title, output_payload()["document_title"])
        self.assertEqual(len(calls), 2)

    def test_document_is_readable_and_uses_tables(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scenario.docx"
            build_scenario_plan_document(
                INPUTS,
                output_path=path,
                ai_generator=lambda _: json.dumps(output_payload(), ensure_ascii=False),
            )
            with ZipFile(path) as archive:
                self.assertIn("word/document.xml", archive.namelist())
            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            self.assertIn("Подготовка", text)
            self.assertIn("Логистика", text)
            self.assertNotIn("Описание идеи:", text)
            self.assertNotIn("Что проверить вручную", text)
            self.assertGreaterEqual(len(document.tables), 5)

    def test_second_invalid_output_returns_controlled_error(self):
        with self.assertRaises(ScenarioPlanGenerationError):
            build_scenario_plan_document(INPUTS, ai_generator=lambda _: "{}")

    def test_legacy_draft_keys_migrate_without_using_description_as_title(self):
        normalized = normalize_inputs(
            "scenario-plan",
            {
                "description": "Старая идея мероприятия",
                "duration": "3 дня",
                "participants": "20 участников",
                "details": "Нужен проектор",
            },
        )
        self.assertEqual(normalized["event_idea"], "Старая идея мероприятия")
        self.assertEqual(normalized["schedule"], "3 дня")
        self.assertEqual(normalized["team_equipment_constraints"], "Нужен проектор")
        self.assertNotIn("event_title", normalized)

    def test_failed_generation_does_not_spend_free_attempt(self):
        client = TestClient(app)
        before = client.get("/api/usage").json()["modules"]["scenario-plan"]["free_attempt_available"]
        with patch(
            "app.services.module_engine.build_scenario_plan_document",
            side_effect=ScenarioPlanGenerationError(USER_FRIENDLY_SCENARIO_ERROR),
        ):
            response = client.post(
                "/api/module-runs",
                json={"module_slug": "scenario-plan", "contest_slug": "pfki", "inputs": INPUTS},
            )
        after = client.get("/api/usage").json()["modules"]["scenario-plan"]["free_attempt_available"]
        self.assertEqual(response.status_code, 400)
        self.assertTrue(before)
        self.assertTrue(after)


if __name__ == "__main__":
    unittest.main()
